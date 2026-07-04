import os
import json
import base64
import tempfile
import uuid
import re
import threading
import mimetypes
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, session

app = Flask(__name__)
# Change this for your own deployment
app.secret_key = 'your-secret-key-here'

UPLOAD_FOLDER = tempfile.mkdtemp()

# ---------------------------------------------------------
# ORGANIZATION MAPPING: Update these rows for your own Excel
# ---------------------------------------------------------
DISTRICTS = {
    'REGION_ALPHA': 4,
    'REGION_BETA':  19,
    'REGION_GAMMA': 34,
    'REGION_DELTA': 49,
    'REGION_OMEGA': 64,
}

SERVICE_COLS = {
    'sunday':      4,
    'monday':      14,
    'thursday':    24,
    'home_caring': 34,
}

SERVICE_LABELS = {
    'sunday':      'Sunday Worship Service',
    'monday':      'Monday Bible Study',
    'thursday':    'Thursday Revival Service',
    'home_caring': 'Home Caring Fellowship',
}

ROW_OFFSETS = {
    'adult_men':       2,
    'adult_women':     3,
    'campus_male':     4,
    'campus_female':   5,
    'youth_boys':      6,
    'youth_girls':     7,
    'children_boys':   8,
    'children_girls':  9,
    'visitors_male':   10,
    'visitors_female': 11,
}

CATEGORY_LABELS = {
    'adult_men':       'Adult Men',
    'adult_women':     'Adult Women',
    'campus_male':     'Campus Male',
    'campus_female':   'Campus Female',
    'youth_boys':      'Youth Boys',
    'youth_girls':     'Youth Girls',
    'children_boys':   'Children Boys',
    'children_girls':  'Children Girls',
    'visitors_male':   'Visitors Male',
    'visitors_female': 'Visitors Female',
}

_store = {}
_store_lock = threading.Lock()

EXTRACTION_PROMPT = """You are processing an attendance record for an organization.

Extract ALL attendance numbers from this document. Return ONLY a valid JSON object — no markdown, no extra text.

Look for:
- Location/District name (e.g. REGION_ALPHA, REGION_BETA, REGION_GAMMA, REGION_DELTA, REGION_OMEGA — or whatever is written)
- Month and year
- Four possible services: Sunday Worship Service, Monday Bible Study, Thursday Revival Service, Home Caring Fellowship
- For each service: up to 5 weeks of attendance data
- Categories per service: Adult (Men/Male + Women/Female), Campus (Male + Female), Youth (Boys/Male + Girls/Female), Children (Boys/Male + Girls/Female), Visitors (Male + Female)

RULES:
- Use integers for numbers that are clearly recorded
- Use null for blank cells, dashes, or truly missing data
- Use 0 ONLY if the document explicitly writes "0"
- Do NOT include totals rows, averages, offering/tithe amounts
- If a service has no data at all, use [null,null,null,null,null] for every category
- Week 5 is often blank — use null unless filled in

Return EXACTLY this JSON structure:
{
  "district": "REGION NAME IN CAPS",
  "month": "MONTH",
  "year": "YEAR",
  "services": {
    "sunday": {
      "adult_men":      [w1, w2, w3, w4, w5],
      "adult_women":    [w1, w2, w3, w4, w5],
      "campus_male":    [w1, w2, w3, w4, w5],
      "campus_female":  [w1, w2, w3, w4, w5],
      "youth_boys":     [w1, w2, w3, w4, w5],
      "youth_girls":    [w1, w2, w3, w4, w5],
      "children_boys":  [w1, w2, w3, w4, w5],
      "children_girls": [w1, w2, w3, w4, w5],
      "visitors_male":  [w1, w2, w3, w4, w5],
      "visitors_female":[w1, w2, w3, w4, w5]
    },
    "monday":      {},
    "thursday":    {},
    "home_caring": {}
  }
}"""


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_store_id():
    if 'store_id' not in session:
        session['store_id'] = uuid.uuid4().hex
    return session['store_id']


def get_store():
    sid = get_store_id()
    with _store_lock:
        if sid not in _store:
            _store[sid] = {'template_path': None, 'extracted': {}, 'output_path': None}
        return _store[sid]


def extract_text_from_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        return '\n'.join(parts)
    except Exception as e:
        return f"Error reading DOCX: {e}"


def match_district(name):
    name = name.upper().strip()
    for known in DISTRICTS:
        if known in name or name in known:
            return known
    return name


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html',
                           districts=list(DISTRICTS.keys()),
                           service_labels=SERVICE_LABELS,
                           category_labels=CATEGORY_LABELS)


@app.route('/api/set-key', methods=['POST'])
def set_key():
    key = request.json.get('api_key', '').strip()
    if not key:
        return jsonify({'success': False, 'error': 'Please enter an API key'})
    session['api_key'] = key
    return jsonify({'success': True})


@app.route('/api/check-key', methods=['GET'])
def check_key():
    return jsonify({'has_key': 'api_key' in session})


@app.route('/api/upload-template', methods=['POST'])
def upload_template():
    if 'template' not in request.files:
        return jsonify({'success': False, 'error': 'No file received'})
    f = request.files['template']
    if not f.filename.lower().endswith('.xlsx'):
        return jsonify({'success': False, 'error': 'Only .xlsx files accepted'})

    fname = f"template_{uuid.uuid4().hex}.xlsx"
    fpath = os.path.join(UPLOAD_FOLDER, fname)
    f.save(fpath)

    try:
        from openpyxl import load_workbook
        wb = load_workbook(fpath)
        if 'DATA INPUT' not in wb.sheetnames:
            return jsonify({'success': False, 'error': 'Template missing "DATA INPUT" sheet'})
        ws = wb['DATA INPUT']
        found = []
        for dname, hrow in DISTRICTS.items():
            val = ws.cell(row=hrow, column=1).value
            if val and dname in str(val).upper():
                found.append(dname)
        wb.close()

        store = get_store()
        store['template_path'] = fpath
        store['extracted'] = {}

        return jsonify({'success': True, 'districts': found, 'filename': f.filename})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/process-document', methods=['POST'])
def process_document():
    if 'api_key' not in session:
        return jsonify({'success': False, 'error': 'API key not set'})

    if 'document' not in request.files:
        return jsonify({'success': False, 'error': 'No document received'})

    f = request.files['document']
    
    # 1. Try to get extension from filename
    ext = Path(f.filename).suffix.lower()
    
    # 2. Clipboard Fallback: Deduce extension from MIME type if pasted
    if not ext and f.content_type:
        ext = mimetypes.guess_extension(f.content_type) or ''
        if ext == '.jpe':
            ext = '.jpeg'

    # 3. Validate File Type
    allowed_exts = {'.pdf', '.png', '.jpg', '.jpeg', '.webp', '.docx'}
    if ext not in allowed_exts:
        return jsonify({'success': False, 'error': f'Unsupported file type (got {ext or f.content_type}). Please use PDF, DOCX, or Images.'})

    tmp = os.path.join(UPLOAD_FOLDER, f"doc_{uuid.uuid4().hex}{ext}")
    f.save(tmp)

    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=session['api_key'])

        if ext in ('.jpg', '.jpeg', '.png', '.webp', '.pdf'):
            mime = 'application/pdf' if ext == '.pdf' else (
                   'image/webp' if ext == '.webp' else (
                   'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/png'))
            
            with open(tmp, 'rb') as file_data:
                part = types.Part.from_bytes(
                    data=file_data.read(),
                    mime_type=mime,
                )

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    EXTRACTION_PROMPT,
                    part
                ],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json"
                )
            )
            raw = response.text

        elif ext == '.docx':
            text = extract_text_from_docx(tmp)
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    f"{EXTRACTION_PROMPT}\n\nDocument text:\n{text}"
                ],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json"
                )
            )
            raw = response.text

        if not raw:
            raise Exception("Received empty response from the model.")

        raw = raw.strip()
        raw = re.sub(r'^```(?:json)?\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        extracted = json.loads(raw)

        district = match_district(extracted.get('district', ''))
        extracted['district'] = district
        extracted['source_file'] = f.filename

        store = get_store()
        store['extracted'][district] = extracted

        os.remove(tmp)
        return jsonify({'success': True, 'data': extracted,
                        'category_labels': CATEGORY_LABELS,
                        'service_labels': SERVICE_LABELS})

    except json.JSONDecodeError as e:
        if os.path.exists(tmp): os.remove(tmp)
        return jsonify({'success': False, 'error': f'AI returned unparseable data: {e}'})
    except Exception as e:
        if os.path.exists(tmp): os.remove(tmp)
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/update-district', methods=['POST'])
def update_district():
    payload = request.json
    district = payload.get('district')
    data = payload.get('data')
    store = get_store()
    if district and data:
        store['extracted'][district] = data
        return jsonify({'success': True})
    return jsonify({'success': False, 'error': 'Missing district or data'})


@app.route('/api/get-extracted', methods=['GET'])
def get_extracted():
    store = get_store()
    return jsonify({
        'extracted': store['extracted'],
        'category_labels': CATEGORY_LABELS,
        'service_labels': SERVICE_LABELS
    })


@app.route('/api/write', methods=['POST'])
def write():
    store = get_store()
    if not store['template_path']:
        return jsonify({'success': False, 'error': 'No template loaded'})
    if not store['extracted']:
        return jsonify({'success': False, 'error': 'No data to write'})

    edits = request.json.get('data', {}) if request.json else {}
    for d, v in edits.items():
        store['extracted'][d] = v

    try:
        from openpyxl import load_workbook
        wb = load_workbook(store['template_path'])
        ws = wb['DATA INPUT']

        summary = []
        total_cells = 0
        total_written = 0

        for dname, ddata in store['extracted'].items():
            if dname not in DISTRICTS:
                continue
            hrow = DISTRICTS[dname]
            dist_summary = {
                'district': dname,
                'services': {},
                'blanks': [],
                'written': 0,
            }

            services = ddata.get('services', {})
            for svc, col_start in SERVICE_COLS.items():
                svc_data = services.get(svc, {})
                svc_written = 0
                svc_blank_flags = []

                for cat, offset in ROW_OFFSETS.items():
                    row = hrow + offset
                    vals = svc_data.get(cat, [None] * 5)
                    for wi in range(5):
                        val = vals[wi] if wi < len(vals) else None
                        col = col_start + wi
                        cell = ws.cell(row=row, column=col)
                        total_cells += 1

                        if val is not None and not cell.protection.locked:
                            cell.value = int(val)
                            svc_written += 1
                            total_written += 1
                        elif val is None and wi < 4:
                            svc_blank_flags.append(
                                f"{CATEGORY_LABELS[cat]} — W{wi+1}"
                            )

                dist_summary['services'][svc] = {
                    'label': SERVICE_LABELS[svc],
                    'written': svc_written,
                    'blanks': svc_blank_flags,
                }
                dist_summary['written'] += svc_written
                dist_summary['blanks'].extend(
                    [f"{SERVICE_LABELS[svc]}: {b}" for b in svc_blank_flags]
                )

            summary.append(dist_summary)

        out_name = f"ATTENDANCE_FILLED_{uuid.uuid4().hex[:8]}.xlsx"
        out_path = os.path.join(UPLOAD_FOLDER, out_name)
        wb.save(out_path)
        wb.close()

        store['output_path'] = out_path
        store['output_name'] = out_name

        return jsonify({
            'success': True,
            'summary': summary,
            'total_written': total_written,
            'output_name': out_name,
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/download')
def download():
    store = get_store()
    if not store.get('output_path'):
        return 'No file ready', 404
    return send_file(
        store['output_path'],
        as_attachment=True,
        download_name=store.get('output_name', 'attendance.xlsx'),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


if __name__ == '__main__':
    print("\n  ╔══════════════════════════════════════════╗")
    print("  ║      AI Attendance Manager Running       ║")
    print("  ║      Open http://localhost:5000          ║")
    print("  ╚══════════════════════════════════════════╝\n")
    app.run(debug=True, port=5000)